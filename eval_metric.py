'''
@Author: shlll
@Date: 2020-05-15 12:02:30
@License: MIT License
@Email: shlll7347@gmail.com
@Modified by: shlll
@Last Modified time: 2020-05-15 12:02:36
@Description:
'''

import json
import shutil
from pathlib import Path
import json
import warnings
from collections import Counter
import argparse
import os
import subprocess

from tqdm.auto import tqdm
import numpy as np
from PIL import Image, ImageDraw
from docx import Document


class EvalMetirc(object):

    def __init__(self, config_path):
        super().__init__()
        self._read_config(config_path)
        self._preprocess_dataset()
        status = self._export_model()
        if status:
            print('=> Export model not sucess.')
            return
        self.run()

    def run(self):
        for model_name in self.tasks['models']:
            if not model_name in self.models:
                warnings.warn(
                    f'model:"{model_name}" not define in config file.')
                continue
            model_config = self.models[model_name]

            for frozen in model_config['frozens']:
                frozen_path = frozen['path']
                iters = frozen['iters']
                print(f"=> Evaluating model:{model_name}, iter: {iters}")

                # try:
                if model_config['type'] == 'deeplab':
                    from deeplab.eval_teeth import DeepLabModel
                    model = DeepLabModel(frozen_path)
                elif model_config['type'] in ['hrnet', 'hrnetocr']:
                    from hrnet.tools.eval_teeth import HRNetModel
                    model = HRNetModel(
                        frozen_path, frozen['cfg'], model_config['type'])
                else:
                    raise Exception("model type not supported.")

            # except FileNotFoundError:
                # warnings.warn(f'frozen:"{frozen_path}" not exists.')
                # continue
            # else:
                for dataset_name in self.tasks['datasets']:
                    print(f'==> Evaluating dataset: {dataset_name}')
                    if not dataset_name in self.datasets:
                        warnings.warn(
                            f'dataset:"{dataset_name}" not define in config file.')
                        continue
                    dataset_config = self.datasets[dataset_name]
                    datset_base_path = Path(dataset_config['path'])
                    gt_path = dataset_config['gt']
                    image_dir = datset_base_path / 'JPEGImages'
                    out_dir = datset_base_path / 'result'
                    if out_dir.exists():
                        shutil.rmtree(out_dir)
                    out_dir.mkdir()
                    model.run_dir(image_dir, out_dir)
                    self._eval_docx(datset_base_path, gt_path,
                                    dataset_name, model_name, iters)

    def _eval_docx(self, base_path, gt, dataset_name, model_name, iters, tumb_size=(100, 100)):
        doc_path = Path('docs/')
        doc_path.mkdir(exist_ok=True)

        base_path = Path(base_path)

        eval_gt_path = base_path / gt
        ground_truth = base_path / 'SegmentationClass'
        source_img = base_path / 'JPEGImages'
        source_img_tumb = base_path / 'JPEGImages_tumb'
        # gt_vis = base_path / 'newSegmentationClass_vis'
        gt_vis = base_path / 'SegmentationClass'
        gt_vis_tumb = base_path / 'newSegmentationClass_vis_tumb'

        numpy_result = base_path / 'result'
        vis_result = base_path / 'result_vis'
        vis_result_tumb = base_path / 'result_vis_tumb'

        if vis_result.exists():
            shutil.rmtree(vis_result)
        if vis_result_tumb.exists():
            shutil.rmtree(vis_result_tumb)

        vis_result.mkdir(exist_ok=True)
        vis_result_tumb.mkdir(exist_ok=True)
        gt_vis_tumb.mkdir(exist_ok=True)
        source_img_tumb.mkdir(exist_ok=True)

        for result_np_path in tqdm(list(numpy_result.glob('*.npy'))):
            # Get file paths.
            source_img_path = next(source_img.glob(result_np_path.stem + '.*'))
            source_img_tumb_path = source_img_tumb / source_img_path.name
            gt_vis_path = gt_vis / (result_np_path.stem + '.png')
            gt_vis_tumb_path = gt_vis_tumb / (result_np_path.stem + '.png')
            res_vis_path = vis_result / (result_np_path.stem + '.jpg')
            res_vis_tumb_path = vis_result_tumb / \
                (result_np_path.stem + '.jpg')

            # Open images and numpy arrays.
            res_np = np.load(str(result_np_path))
            src_img = Image.open(str(source_img_path)).resize(res_np.shape)
            gt_vis_img = np.array(Image.open(gt_vis_path).resize(res_np.shape))
            gt_vis_img[gt_vis_img == 1] = 128
            gt_vis_img[gt_vis_img == 2] = 255
            gt_vis_img = Image.fromarray(gt_vis_img)

            # Draw result images.
            dst_img = src_img.copy()
            draw_img = ImageDraw.Draw(dst_img, mode='RGBA')
            x, y = np.where(res_np == 2)
            dental_point = np.vstack((y, x)).T.flatten()
            draw_img.point(list(dental_point), fill=(255, 255, 0, 64))
            # Save result images.
            dst_img.resize((513, 513)).save(res_vis_path)
            gt_vis_img.resize(tumb_size).save(gt_vis_tumb_path)
            dst_img.resize(tumb_size).save(res_vis_tumb_path)
            src_img.resize(tumb_size).save(source_img_tumb_path)

        document = Document()
        document.add_heading('牙菌斑测试结果', 0)
        paragraph = document.add_paragraph()
        doc_rows = len(list(source_img.glob('*'))) + 1
        table = document.add_table(rows=doc_rows, cols=6)

        iou_sum = 0
        acc_sum = 0
        count = 0

        for source_img_path in tqdm(sorted(list(source_img.glob('*')))):
            key = source_img_path.stem
            source_img_tumb_path = source_img_tumb / source_img_path.name
            gt_img_path = eval_gt_path / (key + '.png')
            gt_vis_tumb_path = gt_vis_tumb / (key + '.png')
            teeth_res_vis_path = vis_result_tumb / (key + '.jpg')
            teeth_numpy_path = numpy_result / (key + '.npy')

            # if np.load(teeth_numpy_path).shape != (513, 513):
            #     continue

            iou, acc = self._compute_metric(teeth_numpy_path, gt_img_path)

            count += 1
            acc_sum += acc
            iou_sum += iou
            iou = f'{iou:.5f}'
            acc = f'{acc:.5f}'

            row = table.rows[count]
            self._docx_add_row(
                row, [source_img_tumb_path, gt_vis_tumb_path, teeth_res_vis_path], iou, acc)

        miou = iou_sum / count
        macc = acc_sum / count
        paragraph.add_run(f"MIOU:{miou}, MACC:{macc}")
        document.save(
            str(doc_path / f'{model_name}_{dataset_name}_{iters}.docx'))

    def _preprocess_dataset(self):
        for name, dataset in self.datasets.items():
            if "ori_path" in dataset:
                print(f"==>数据集[{name}]预处理，生成测试集")
                ori_path = Path(dataset['ori_path'])
                ori_src_path = ori_path / 'JPEGImages'
                ori_gt_path = ori_path / 'SegmentationClass'
                dst_path = Path(dataset['path'])
                dst_src_path = dst_path / 'JPEGImages'
                dst_gt_path = dst_path / dataset['gt']

                if dst_path.exists():
                    shutil.rmtree(dst_path)
                    # continue
                dst_path.mkdir()
                dst_src_path.mkdir()
                dst_gt_path.mkdir()

                val_file = ori_path / "ImageSets/Segmentation/val.txt"
                val_file_list = self._read_summary_file(val_file)
                for val_file_name in tqdm(val_file_list):
                    val_file_src_name = val_file_name + '.jpg'
                    val_file_gt_name = val_file_name + '.png'
                    ori_src_img = ori_src_path / val_file_src_name
                    dst_src_img = dst_src_path / val_file_src_name
                    ori_gt_img = ori_gt_path / val_file_gt_name
                    dst_gt_img = dst_gt_path / val_file_gt_name

                    if ori_src_img.exists() and ori_gt_img.exists():
                        os.link(ori_src_img, dst_src_img)
                        os.link(ori_gt_img, dst_gt_img)

    def _read_summary_file(self, path):
        summary_list = []
        with open(path, 'r') as f:
            for line in f:
                if (line[-1] == '\n'):
                    line = line[:-1]
                summary_list.append(line)

        return summary_list

    def _export_model(self):
        for name, model in self.models.items():
            if model['type'] == 'deeplab':
                export = model['export']
                override = export['override']

                for iters in export["iters"]:
                    frozen_path = Path(export['dir']) / f'{name}_{iters}.pb'
                    ckpt_path = f'train_log_{name}/model.ckpt-{iters}'
                    if frozen_path.exists() and not override:
                        continue
                    cmd = f'''python deeplab/export_model.py \
                                --atrous_rates=6 \
                                --atrous_rates=12 \
                                --atrous_rates=18 \
                                --output_stride=16 \
                                --decoder_output_stride=4 \
                                --model_variant="xception_65" \
                                --num_classes={export['num_classes']} \
                                --checkpoint_path="{ckpt_path}" \
                                --export_path="{str(frozen_path)}"
                            '''
                    ex = subprocess.Popen(
                        cmd, stdout=subprocess.PIPE, shell=True)
                    out, err = ex.communicate()
                    status = ex.wait()

                    if not status:
                        return status
            elif model['type'] in ['hrnet', 'hrnetocr']:
                import torch
                frozens = model['frozens']
                export_dir = model['export']['dir']
                for frozen in frozens:
                    model_path = frozen['path']
                    frozen['iters'] = torch.load(model_path)['epoch']
                    shutil.copy(model_path, Path(export_dir) /
                                f"{name}_{frozen['iters']}.pth.tar")

        return 0

    def _read_config(self, path):
        """Read json config file.

        Arguments:
            path {str} -- path to the json file.

        Raises:
            FileNotFoundError: The path not exists.
        """
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f'{path} not found.')

        with open(path, 'r') as f:
            config = json.load(f)

        self.datasets = {}
        self.models = {}
        self.tasks = config['enabled']

        for dataset in config['datasets']:
            option = {
                'path': dataset['base'],
                'gt': dataset['gt']
            }

            if 'ori_path' in dataset:
                option['ori_path'] = dataset['ori_path']

            self.datasets[dataset['name']] = option
        for model in config['models']:
            self.models[model['name']] = {
                'type': model['type'],
                'export': model['export'],
                'frozens': model['frozens']
            }

    def _compute_metric(self, pred_path, gt_path):
        res_np = np.load(pred_path).astype('uint8')
        gt_img = Image.open(gt_path).resize(res_np.shape[-2:])
        gt_np = np.asarray(gt_img, dtype='uint8').copy()
        gt_np[gt_np == 125] = 1
        gt_np[gt_np == 255] = 2

        res_np = res_np * (gt_np > 0)
        intersection = res_np * (res_np == gt_np)

        area_pred = Counter(res_np.flatten())
        area_lab = Counter(gt_np.flatten())
        area_intersection = Counter(intersection.flatten())
        area_union = area_pred + area_lab - area_intersection

        iou = self._sum_iou(area_intersection) / self._sum_iou(area_union)

        acc = (res_np.flatten() == gt_np.flatten()
               ).sum() / res_np.flatten().size

        return iou, acc

    def _sum_iou(self, x):
        return x[1]+x[2]

    def _docx_add_row(self, row, pics, iou, acc):
        for i, pic in enumerate(pics):
            row.cells[i].paragraphs[0].add_run().add_picture(str(pic))
        row.cells[i+1].text = iou
        row.cells[i+2].text = acc


if __name__ == "__main__":
    paser = argparse.ArgumentParser(description="Eval metrics.")
    paser.add_argument("-c", "--config", type=str,
                       required=True, help="The configure directory.")
    args = paser.parse_args()

    EvalMetirc(args.config)
