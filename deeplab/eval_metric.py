'''
@Author: shlll
@Date: 2020-05-15 12:02:30
@License: MIT License
@Email: shlll7347@gmail.com
@Modified by: shlll
@Last Modified time: 2020-05-15 12:02:36
@Description:
'''

import json
from pathlib import Path
import json
import warnings
from collections import Counter
import argparse
import os
import subprocess

from tqdm.auto import tqdm
import numpy as np
from PIL import Image, ImageDraw
from docx import Document

from deeplab.eval_teeth import DeepLabModel


class EvalMetirc(object):

    def __init__(self, config_path):
        super().__init__()
        self._read_config(config_path)
        status = self._export_model()
        if status:
            print('=> Export model not sucess.')
            return
        self.run()

    def run(self):
        for model_name in self.tasks['models']:
            if not model_name in self.models:
                warnings.warn(
                    f'model:"{model_name}" not define in config file.')
                continue
            model_config = self.models[model_name]

            for frozen in model_config['frozens']:
                frozen_path = frozen['path']
                iters = frozen['iters']
                print(f"=> Evaluating model:{model_name}, iter: {iters}")

                try:
                    model = DeepLabModel(frozen_path)
                except FileNotFoundError:
                    warnings.warn(f'frozen:"{frozen_path}" not exists.')
                    continue
                else:
                    for dataset_name in self.tasks['datasets']:
                        print(f'==> Evaluating dataset: {dataset_name}')
                        if not dataset_name in self.datasets:
                            warnings.warn(
                                f'dataset:"{dataset_name}" not define in config file.')
                            continue
                        dataset_config = self.datasets[dataset_name]
                        datset_base_path = Path(dataset_config['path'])
                        gt_path = dataset_config['gt']
                        image_dir = datset_base_path / 'JPEGImages'
                        out_dir = datset_base_path / 'result'
                        model.run_dir(image_dir, out_dir)
                        self._eval_docx(datset_base_path, gt_path,
                                        dataset_name, model_name, iters)

    def _eval_docx(self, base_path, gt, dataset_name, model_name, iters, tumb_size=(100, 100)):
        doc_path = Path('docs/')
        doc_path.mkdir(exist_ok=True)

        base_path = Path(base_path)

        eval_gt_path = base_path / gt
        ground_truth = base_path / 'SegmentationClass'
        source_img = base_path / 'JPEGImages'
        source_img_tumb = base_path / 'JPEGImages_tumb'
        gt_vis = base_path / 'newSegmentationClass_vis'
        gt_vis_tumb = base_path / 'newSegmentationClass_vis_tumb'

        numpy_result = base_path / 'result'
        vis_result = base_path / 'result_vis'
        vis_result_tumb = base_path / 'result_vis_tumb'

        vis_result.mkdir(exist_ok=True)
        vis_result_tumb.mkdir(exist_ok=True)
        gt_vis_tumb.mkdir(exist_ok=True)
        source_img_tumb.mkdir(exist_ok=True)

        for result_np_path in tqdm(list(numpy_result.glob('*.npy'))):
            # Get file paths.
            source_img_path = next(source_img.glob(result_np_path.stem + '.*'))
            source_img_tumb_path = source_img_tumb / source_img_path.name
            gt_vis_path = gt_vis / (result_np_path.stem + '.png')
            gt_vis_tumb_path = gt_vis_tumb / (result_np_path.stem + '.png')
            res_vis_path = vis_result / (result_np_path.stem + '.jpg')
            res_vis_tumb_path = vis_result_tumb / \
                (result_np_path.stem + '.jpg')

            # Open images and numpy arrays.
            src_img = Image.open(str(source_img_path)).resize((513, 513))
            gt_vis_img = Image.open(gt_vis_path)
            res_np = np.load(str(result_np_path))

            # Draw result images.
            dst_img = src_img.copy()
            draw_img = ImageDraw.Draw(dst_img, mode='RGBA')
            x, y = np.where(res_np == 2)
            dental_point = np.vstack((y, x)).T.flatten()
            draw_img.point(list(dental_point), fill=(255, 255, 0, 64))
            # Save result images.
            dst_img.save(res_vis_path)
            gt_vis_img.resize(tumb_size).save(gt_vis_tumb_path)
            dst_img.resize(tumb_size).save(res_vis_tumb_path)
            src_img.resize(tumb_size).save(source_img_tumb_path)

        document = Document()
        document.add_heading('牙菌斑测试结果', 0)
        paragraph = document.add_paragraph()
        doc_rows = len(list(source_img.glob('*'))) + 1
        table = document.add_table(rows=doc_rows, cols=6)

        iou_sum = 0
        acc_sum = 0
        count = 0

        for i, source_img_path in enumerate(tqdm(sorted(list(source_img.glob('*'))))):
            key = source_img_path.stem
            source_img_tumb_path = source_img_tumb / source_img_path.name
            gt_img_path = eval_gt_path / (key + '.png')
            gt_vis_tumb_path = gt_vis_tumb / (key + '.png')
            teeth_res_vis_path = vis_result_tumb / (key + '.jpg')
            teeth_numpy_path = numpy_result / (key + '.npy')
            iou, acc = self._compute_metric(teeth_numpy_path, gt_img_path)

            count += 1
            acc_sum += acc
            iou_sum += iou
            iou = f'{iou:.5f}'
            acc = f'{acc:.5f}'

            row = table.rows[i+1]
            self._docx_add_row(
                row, [source_img_tumb_path, gt_vis_tumb_path, teeth_res_vis_path], iou, acc)

        miou = iou_sum / count
        macc = acc_sum / count
        paragraph.add_run(f"MIOU:{miou}, MACC:{macc}")
        document.save(
            str(doc_path / f'{model_name}_{dataset_name}_{iters}.docx'))

    def _export_model(self):
        for name, model in self.models.items():
            export = model['export']
            override = export['override']

            for iters in export["iters"]:
                frozen_path = Path(export['dir']) / f'{name}_{iters}.pb'
                ckpt_path = f'train_log_{name}/model.ckpt-{iters}'
                if frozen_path.exists() and not override:
                    continue
                cmd = f'''python deeplab/export_model.py \
                            --atrous_rates=6 \
                            --atrous_rates=12 \
                            --atrous_rates=18 \
                            --output_stride=16 \
                            --decoder_output_stride=4 \
                            --model_variant="xception_65" \
                            --num_classes={export['num_classes']} \
                            --checkpoint_path="{ckpt_path}" \
                            --export_path="{str(frozen_path)}"
                        '''
                ex = subprocess.Popen(cmd, stdout=subprocess.PIPE, shell=True)
                out, err = ex.communicate()
                status = ex.wait()

                if not status:
                    return status

        return 0

    def _read_config(self, path):
        """Read json config file.

        Arguments:
            path {str} -- path to the json file.

        Raises:
            FileNotFoundError: The path not exists.
        """
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f'{path} not found.')

        with open(path, 'r') as f:
            config = json.load(f)

        self.datasets = {}
        self.models = {}
        self.tasks = config['enabled']

        for dataset in config['datasets']:
            self.datasets[dataset['name']] = {
                'path': dataset['base'],
                'gt': dataset['gt']
            }
        for model in config['models']:
            self.models[model['name']] = {
                'export': model['export'],
                'frozens': model['frozens']
            }

    def _compute_metric(self, pred_path, gt_path):
        res_np = np.load(pred_path).astype('uint8')
        gt_img = Image.open(gt_path).resize(res_np.shape[-2:])
        gt_np = np.asarray(gt_img, dtype='uint8').copy()
        gt_np[gt_np == 125] = 1
        gt_np[gt_np == 255] = 2

        res_np = res_np * (gt_np > 0)
        intersection = res_np * (res_np == gt_np)

        area_pred = Counter(res_np.flatten())
        area_lab = Counter(gt_np.flatten())
        area_intersection = Counter(intersection.flatten())
        area_union = area_pred + area_lab - area_intersection

        iou = self._sum_iou(area_intersection) / self._sum_iou(area_union)

        acc = (res_np.flatten() == gt_np.flatten()
               ).sum() / res_np.flatten().size

        return iou, acc

    def _sum_iou(self, x):
        return x[1]+x[2]

    def _docx_add_row(self, row, pics, iou, acc):
        for i, pic in enumerate(pics):
            row.cells[i].paragraphs[0].add_run().add_picture(str(pic))
        row.cells[i+1].text = iou
        row.cells[i+2].text = acc


if __name__ == "__main__":
    paser = argparse.ArgumentParser(description="Eval metrics.")
    paser.add_argument("-c", "--config", type=str,
                       required=True, help="The configure directory.")
    args = paser.parse_args()

    EvalMetirc(args.config)
